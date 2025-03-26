from typing import List
from docx import Document
import pandas as pd
from glob import glob
import configs.general.config as c


def join_text(counter: int, all_text: Document, stop_condition: str | List[str]) -> (int, List[str]):
    counter += 1
    text_list = []
    if type(stop_condition) is list:
        stop_condition = tuple(stop_condition)
    while counter < len(all_text.paragraphs) and not all_text.paragraphs[counter].text.strip().startswith(stop_condition):
        text_list.append(all_text.paragraphs[counter].text)
        counter += 1

    return counter - 1, '\n'.join(text_list)


text_df = pd.DataFrame()
for doc_path in glob('data/sal_parse/*docx'):
    doc = Document(doc_path)
    # sub_paragraph_df = pd.DataFrame(
    #     columns=['Name', 'Main_Sal', 'Description', 'Min_conditions', 'Additional_conditions', 'Operating_conditions',
    #              'Consideration_clauses_for_adjustment'])

    sub_paragraph_list = []
    sub_paragraph_dict = {}
    i = 0
    while i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()
        # if text.startswith('סל'):
        #     main_name = text.replace(":", "")
        #     sub_paragraph_dict['Main_Sal'] = main_name
        if text.startswith('תת סל'):
            # if 'Main_Sal' not in sub_paragraph_dict:
            #     sub_paragraph_dict['Main_Sal'] = main_name

            print(len(sub_paragraph_list)+1, '\t', doc.paragraphs[i].text)
            sub_paragraph_dict['Name'] = text.replace("תת סל: ", "")
        if text.startswith('תיאור תת הסל'):
            i, sub_paragraph_dict['Description'] = join_text(counter=i, all_text=doc, stop_condition='תנאי סף')

        if text.startswith('תנאי סף'):
            i, sub_paragraph_dict['Min_conditions'] = join_text(counter=i, all_text=doc, stop_condition=['תנאים נוספים', 'תת סל'])

        # if text.startswith('תנאים נוספים'):
        #     i, sub_paragraph_dict['Additional_conditions'] = join_text(counter=i, all_text=doc, stop_condition='תנאי הפעלה')
        #
        # if text.startswith('תנאי הפעלה'):
        #     i, sub_paragraph_dict['Operating_conditions'] = join_text(counter=i, all_text=doc,
        #                                                               stop_condition='בעת בחינת הצוות המקצועי')
        #
        # if text.startswith('בעת בחינת הצוות המקצועי'):
        #     i, sub_paragraph_dict['Consideration_clauses_for_adjustment'] = join_text(counter=i, all_text=doc,
        #                                                                               stop_condition=['סל', 'תת סל:'])

            sub_paragraph_list.append(sub_paragraph_dict)
            sub_paragraph_dict = {}

        i += 1

    df = pd.DataFrame(sub_paragraph_list)
    text_df = pd.concat([text_df, df], ignore_index=True)

tables = Document('data_old/table_parse/single_table.docx')
data = []
for table in tables.tables:
    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)

        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        for_df = {'Name': row_data['שם'].strip(),
                  'Description': row_data['תיאור התפקיד'].strip() + row_data['נושאי ליבה ופעולות נדרשות'].strip() + row_data['תוצרים אפשריים ותוצאות'].strip(),
                  'Min_conditions': row_data['תנאי סף ופירוט ניסיון'].strip() + row_data['מסמכים נדרשים/פירוט ניסיון'].strip()}

        data.append(for_df)

tables_df = pd.DataFrame(data)

final_df = pd.concat([text_df, tables_df]).reset_index(drop=True)

final_df.to_csv(c.subjects_path, index=False, encoding='utf-8-sig')
